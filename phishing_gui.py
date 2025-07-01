import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import re
import email
import joblib
import numpy as np
import requests
import time
from datetime import datetime
from urllib.parse import urlparse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import hashlib
import base64
import warnings
from sklearn.exceptions import InconsistentVersionWarning
import threading
import email.parser

try:
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    SELENIUM_AVAILABLE = True
except ImportError:
    SELENIUM_AVAILABLE = False

try:
    from PIL import Image, ImageTk
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

warnings.filterwarnings("ignore", category=InconsistentVersionWarning)

MODEL_PATH = os.path.abspath("phishing_model_v2.joblib")

class PhishingDetector:
    def __init__(self):
        self.model_path = MODEL_PATH
        self.components = None
        self.vt_api_key = os.getenv('VIRUSTOTAL_API_KEY')
        self.load_model()

    def load_model(self):
        try:
            if not os.path.exists(self.model_path):
                raise FileNotFoundError(f"Model file not found: {self.model_path}")
            self.components = joblib.load(self.model_path)
        except Exception as e:
            self.components = None
            print(f"Error loading model: {e}")

    def validate_input(self, text):
        if len(text.strip()) == 0:
            raise ValueError("Input cannot be empty")
        if len(text) > 10000:
            raise ValueError("Input text is too long (max 10,000 characters)")
        return text.strip()

    def preprocess_text(self, text):
        text = re.sub(r'http\S+', '', text)
        text = re.sub(r'[^\w\s]', '', text)
        text = text.lower()
        text = re.sub(r'\s+', ' ', text).strip()
        return text

    def extract_urls(self, text):
        url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
        urls = re.findall(url_pattern, text)
        return list(set(urls))

    def validate_url(self, url):
        try:
            parsed = urlparse(url)
            if not parsed.scheme or not parsed.netloc:
                return False
            if parsed.scheme not in ['http', 'https']:
                return False
            return True
        except Exception:
            return False

    def analyze_url_virustotal(self, url):
        if not self.vt_api_key:
            return {"error": "VirusTotal API key not available"}
        try:
            url_id = base64.urlsafe_b64encode(url.encode()).decode().strip("=")
            headers = {"X-Apikey": self.vt_api_key, "Content-Type": "application/json"}
            response = requests.get(
                f"https://www.virustotal.com/api/v3/urls/{url_id}",
                headers=headers, timeout=30)
            if response.status_code == 404:
                submit_response = requests.post(
                    "https://www.virustotal.com/api/v3/urls",
                    headers=headers, data={"url": url}, timeout=30)
                if submit_response.status_code == 200:
                    time.sleep(15)
                    response = requests.get(
                        f"https://www.virustotal.com/api/v3/urls/{url_id}",
                        headers=headers, timeout=30)
            if response.status_code == 200:
                data = response.json()
                stats = data.get('data', {}).get('attributes', {}).get('last_analysis_stats', {})
                malicious = stats.get('malicious', 0)
                suspicious = stats.get('suspicious', 0)
                clean = stats.get('harmless', 0)
                undetected = stats.get('undetected', 0)
                total_scans = malicious + suspicious + clean + undetected
                if total_scans == 0:
                    return {"status": "No analysis data available", "url": url}
                if malicious > 0:
                    classification = "Malicious"
                elif suspicious > 2:
                    classification = "Suspicious"
                else:
                    classification = "Legitimate"
                return {
                    "url": url,
                    "classification": classification,
                    "malicious": malicious,
                    "suspicious": suspicious,
                    "clean": clean,
                    "total_scans": total_scans,
                    "scan_date": data.get('data', {}).get('attributes', {}).get('last_analysis_date')
                }
            else:
                return {"error": f"API request failed with status {response.status_code}"}
        except requests.exceptions.Timeout:
            return {"error": "Request timeout"}
        except requests.exceptions.RequestException as e:
            return {"error": f"Request failed: {str(e)}"}
        except Exception as e:
            return {"error": f"Analysis failed: {str(e)}"}

    def predict_email(self, email_text):
        try:
            email_text = self.validate_input(email_text)
            if not self.components:
                return {"error": "Model not loaded"}
            tfidf_vectorizer = self.components['tfidf_vectorizer']
            voting_classifier = self.components['voting_classifier']
            label_encoder = self.components['label_encoder']
            processed_text = self.preprocess_text(email_text)
            email_tfidf = tfidf_vectorizer.transform([processed_text])
            prediction = voting_classifier.predict(email_tfidf)
            probabilities = voting_classifier.predict_proba(email_tfidf)
            return {
                'class': label_encoder.inverse_transform(prediction)[0],
                'probability': float(np.max(probabilities)),
                'probabilities': {
                    label_encoder.classes_[0]: float(probabilities[0][0]),
                    label_encoder.classes_[1]: float(probabilities[0][1])
                }
            }
        except Exception as e:
            return {"error": f"Prediction failed: {str(e)}"}

def analyze_headers_from_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            msg = email.message_from_file(f)
        parser = email.parser.HeaderParser()
        headers = parser.parsestr(msg.as_string())
        meta = {
            "message-id": "",
            "spf-record": False,
            "dkim-record": False,
            "dmarc-record": False,
            "spoofed": False,
            "ip-address": "",
            "sender-client": "",
            "spoofed-mail": "",
            "dt": "",
            "content-type": "",
            "from": "",
            "reply-to": "",
            "return-path": "",
            "x-originating-ip": "",
            "multiple-received": 0,
            "threat-level": "LEGITIMATE",
            "threat-score": 0
        }
        received_count = 0
        for h in headers.items():
            if h[0].lower() == "message-id":
                meta["message-id"] = h[1]
            if h[0].lower() == "from":
                meta["from"] = h[1]
            if h[0].lower() == "return-path":
                meta["return-path"] = h[1]
            if h[0].lower() == "received":
                received_count += 1
                if received_count == 1:
                    meta["sender-client"] = h[1]
            if h[0].lower() == "x-originating-ip":
                ip = re.search(r"(\d{1,3}\.){3}\d{1,3}", h[1])
                if ip:
                    meta["ip-address"] = str(ip.group())
            if h[0].lower() == "authentication-results":
                if re.search(r"spf=pass", h[1], re.IGNORECASE):
                    meta["spf-record"] = True
                if re.search(r"dkim=pass", h[1], re.IGNORECASE):
                    meta["dkim-record"] = True
                if re.search(r"dmarc=pass", h[1], re.IGNORECASE):
                    meta["dmarc-record"] = True
                if re.search(r"does not designate|spf=fail|dkim=fail|dmarc=fail", h[1], re.IGNORECASE):
                    meta["spoofed"] = True
            if h[0].lower() == "reply-to":
                meta["reply-to"] = h[1]
                meta["spoofed-mail"] = h[1]
            if h[0].lower() == "date":
                meta["dt"] = h[1]
            if h[0].lower() == "content-type":
                meta["content-type"] = h[1]
        meta["multiple-received"] = received_count
        threat_score = 0
        threat_reasons = []
        if meta["spoofed"]:
            threat_score += 40
            threat_reasons.append("Authentication failures detected")
        if not meta["spf-record"]:
            threat_score += 15
            threat_reasons.append("SPF record failed")
        if not meta["dkim-record"]:
            threat_score += 15
            threat_reasons.append("DKIM verification failed")
        if not meta["dmarc-record"]:
            threat_score += 15
            threat_reasons.append("DMARC verification failed")
        if meta["from"] and meta["reply-to"] and meta["from"] != meta["reply-to"]:
            threat_score += 20
            threat_reasons.append("From and Reply-To headers mismatch")
        if meta["multiple-received"] > 8:
            threat_score += 15
            threat_reasons.append("Excessive email routing detected")
        if not meta["message-id"]:
            threat_score += 10
            threat_reasons.append("Missing Message-ID")
        if threat_score >= 70:
            meta["threat-level"] = "MALICIOUS"
        elif threat_score >= 30:
            meta["threat-level"] = "SUSPICIOUS"
        else:
            meta["threat-level"] = "LEGITIMATE"
        meta["threat-score"] = threat_score
        meta["threat-reasons"] = threat_reasons
        return meta
    except Exception as e:
        return {"error": str(e)}

def get_header_recommendation(threat_level, threat_score, threat_reasons):
    if threat_level == "MALICIOUS":
        return "DO NOT INTERACT with this email. Delete immediately and report to IT security team. This email shows strong indicators of malicious intent."
    elif threat_level == "SUSPICIOUS":
        return "EXERCISE CAUTION. Do not click links or download attachments. Verify sender through alternative communication method before taking any action."
    else:
        return "Email appears legitimate. However, always exercise standard email security practices."

#GUI 
class PhishingGUI(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Phishing Email Analyzer")
        self.geometry("900x700")
        self.resizable(True, True)
        self.detector = PhishingDetector()
        self.create_widgets()

    def create_widgets(self):
        tab_control = ttk.Notebook(self)
        self.tab_content = ttk.Frame(tab_control)
        self.tab_headers = ttk.Frame(tab_control)
        tab_control.add(self.tab_content, text='Email Content/URL Analysis')
        tab_control.add(self.tab_headers, text='Header Analysis')
        tab_control.pack(expand=1, fill='both')
        # Content/URL Analysis Tab
        self.create_content_tab()
        # Header Analysis Tab
        self.create_headers_tab()

    def create_content_tab(self):
        frame = self.tab_content
        # Email text input
        lbl = ttk.Label(frame, text="Paste email text or load from file:")
        lbl.pack(anchor='nw', padx=10, pady=5)
        self.email_text = scrolledtext.ScrolledText(frame, height=12)
        self.email_text.pack(fill='x', padx=10, pady=5)
        btn_frame = ttk.Frame(frame)
        btn_frame.pack(fill='x', padx=10, pady=5)
        ttk.Button(btn_frame, text="Load from File", command=self.load_email_file).pack(side='left')
        ttk.Button(btn_frame, text="Analyze", command=self.run_content_analysis).pack(side='left', padx=5)
        self.btn_screenshot = ttk.Button(btn_frame, text="Capture Screenshots", command=self.capture_screenshots, state='disabled')
        self.btn_screenshot.pack(side='left', padx=5)
        self.btn_save_content_log = ttk.Button(btn_frame, text="Save Log", command=self.save_content_log, state='disabled')
        self.btn_save_content_log.pack(side='left', padx=5)
        self.btn_save_content_report = ttk.Button(btn_frame, text="Save Report", command=self.save_content_report, state='disabled')
        self.btn_save_content_report.pack(side='left', padx=5)
        # Results
        ttk.Label(frame, text="Results:").pack(anchor='nw', padx=10, pady=5)
        self.content_results = scrolledtext.ScrolledText(frame, height=10, state='disabled')
        self.content_results.pack(fill='x', expand=False, padx=10, pady=5)
        # Screenshot previews
        self.screenshot_frame = ttk.Frame(frame)
        self.screenshot_frame.pack(fill='both', expand=True, padx=10, pady=5)
        self.screenshot_images = []  # To keep references to PhotoImage objects
        self.content_analysis_results = None
        self.content_report_path = None
        self.content_log_path = None
        self.screenshot_paths = []

    def create_headers_tab(self):
        frame = self.tab_headers
        lbl = ttk.Label(frame, text="Select .eml or header file:")
        lbl.pack(anchor='nw', padx=10, pady=5)
        btn_frame = ttk.Frame(frame)
        btn_frame.pack(fill='x', padx=10, pady=5)
        ttk.Button(btn_frame, text="Load Header File", command=self.load_header_file).pack(side='left')
        ttk.Button(btn_frame, text="Analyze", command=self.run_header_analysis).pack(side='left', padx=5)
        self.btn_save_header_log = ttk.Button(btn_frame, text="Save Log", command=self.save_header_log, state='disabled')
        self.btn_save_header_log.pack(side='left', padx=5)
        self.btn_save_header_report = ttk.Button(btn_frame, text="Save Report", command=self.save_header_report, state='disabled')
        self.btn_save_header_report.pack(side='left', padx=5)
        self.header_file_path = tk.StringVar()
        ttk.Label(frame, textvariable=self.header_file_path).pack(anchor='nw', padx=10, pady=2)
        ttk.Label(frame, text="Results:").pack(anchor='nw', padx=10, pady=5)
        self.header_results = scrolledtext.ScrolledText(frame, height=25, state='disabled')
        self.header_results.pack(fill='both', expand=True, padx=10, pady=5)
        self.header_analysis_results = None
        self.header_report_path = None
        self.header_log_path = None

    # Content/URL Analysis Tab Methods
    def load_email_file(self):
        path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")])
        if path:
            with open(path, 'r', encoding='utf-8') as f:
                self.email_text.delete('1.0', tk.END)
                self.email_text.insert(tk.END, f.read())

    def run_content_analysis(self):
        text = self.email_text.get('1.0', tk.END).strip()
        if not text:
            messagebox.showwarning("Input Required", "Please enter or load email text.")
            return
        self.content_results.config(state='normal')
        self.content_results.delete('1.0', tk.END)
        self.content_results.insert(tk.END, "Analyzing...\n")
        self.content_results.config(state='disabled')
        def worker():
            results = {'email_content': text}
            email_result = self.detector.predict_email(text)
            results['email_analysis'] = email_result
            urls = self.detector.extract_urls(text)
            results['url_analysis'] = []
            if urls:
                for url in urls:
                    if self.detector.validate_url(url):
                        url_result = self.detector.analyze_url_virustotal(url)
                        results['url_analysis'].append(url_result)
                    else:
                        results['url_analysis'].append({"url": url, "error": "Invalid URL"})
            self.content_analysis_results = results
            self.display_content_results(results)
        threading.Thread(target=worker).start()

    def display_content_results(self, results):
        self.content_results.config(state='normal')
        self.content_results.delete('1.0', tk.END)
        email_result = results.get('email_analysis', {})
        self.content_results.insert(tk.END, f"Email Classification: {email_result.get('class', 'N/A')}\n")
        self.content_results.insert(tk.END, f"Confidence: {email_result.get('probability', 0)*100:.2f}%\n")
        self.content_results.insert(tk.END, f"Probabilities: {email_result.get('probabilities', {})}\n\n")
        urls = results.get('url_analysis', [])
        if urls:
            self.content_results.insert(tk.END, f"URLs Found: {len(urls)}\n")
            for i, url_result in enumerate(urls, 1):
                self.content_results.insert(tk.END, f"URL {i}: {url_result.get('url', 'N/A')}\n")
                if 'error' not in url_result:
                    self.content_results.insert(tk.END, f"  Classification: {url_result.get('classification', 'N/A')}\n")
                    self.content_results.insert(tk.END, f"  Malicious: {url_result.get('malicious', 0)}\n")
                    self.content_results.insert(tk.END, f"  Suspicious: {url_result.get('suspicious', 0)}\n")
                    self.content_results.insert(tk.END, f"  Total Scans: {url_result.get('total_scans', 0)}\n")
                else:
                    self.content_results.insert(tk.END, f"  Error: {url_result['error']}\n")
                self.content_results.insert(tk.END, "\n")
            self.btn_screenshot.config(state='normal')
        else:
            self.content_results.insert(tk.END, "No URLs found in email.\n")
            self.btn_screenshot.config(state='disabled')
        self.content_results.config(state='disabled')
        self.btn_save_content_log.config(state='normal')
        self.btn_save_content_report.config(state='normal')
        self.clear_screenshot_previews()

    def clear_screenshot_previews(self):
        for widget in self.screenshot_frame.winfo_children():
            widget.destroy()
        self.screenshot_images = []
        self.screenshot_paths = []

    def capture_screenshots(self):
        if not SELENIUM_AVAILABLE:
            messagebox.showerror("Selenium Not Available", "Selenium and ChromeDriver are required for screenshots.")
            return
        urls = []
        if self.content_analysis_results:
            urls = [u['url'] for u in self.content_analysis_results.get('url_analysis', []) if 'url' in u and 'error' not in u]
        if not urls:
            messagebox.showinfo("No URLs", "No valid URLs to screenshot.")
            return
        self.clear_screenshot_previews()
        def worker():
            screenshots_dir = "screenshots"
            os.makedirs(screenshots_dir, exist_ok=True)
            for url in urls:
                try:
                    chrome_options = Options()
                    chrome_options.add_argument('--headless')
                    chrome_options.add_argument('--no-sandbox')
                    chrome_options.add_argument('--disable-dev-shm-usage')
                    chrome_options.add_argument('--window-size=800,600')
                    driver = webdriver.Chrome(options=chrome_options)
                    driver.set_page_load_timeout(30)
                    driver.get(url)
                    time.sleep(2)
                    url_hash = hashlib.md5(url.encode()).hexdigest()[:8]
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    filename = f"screenshot_{url_hash}_{timestamp}.png"
                    filepath = os.path.join(screenshots_dir, filename)
                    driver.save_screenshot(filepath)
                    driver.quit()
                    self.screenshot_paths.append(filepath)
                    self.show_screenshot_preview(filepath, url)
                except Exception as e:
                    self.show_screenshot_preview(None, url, error=str(e))
        threading.Thread(target=worker).start()

    def show_screenshot_preview(self, filepath, url, error=None):
        def _show():
            frame = ttk.Frame(self.screenshot_frame)
            frame.pack(side='left', padx=5, pady=5)
            if filepath and os.path.exists(filepath):
                if PIL_AVAILABLE:
                    try:
                        img = Image.open(filepath)
                        img.thumbnail((180, 120))
                        photo = ImageTk.PhotoImage(img)
                        lbl_img = tk.Label(frame, image=photo)
                        lbl_img.image = photo  # keep reference
                        lbl_img.pack()
                        self.screenshot_images.append(photo)
                    except Exception as e:
                        ttk.Label(frame, text=f"[Preview error: {e}]").pack()
                else:
                    ttk.Label(frame, text="[PIL not installed for preview]").pack()
            else:
                ttk.Label(frame, text="[Screenshot failed]").pack()
            ttk.Label(frame, text=os.path.basename(filepath) if filepath else url).pack()
            
        self.screenshot_frame.after(0, _show)

    def save_content_log(self):
        if not self.content_analysis_results:
            return
        logs_dir = "logs"
        os.makedirs(logs_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        log_file = os.path.join(logs_dir, f"email_content_log_{timestamp}.txt")
        results = self.content_analysis_results
        with open(log_file, 'w', encoding='utf-8') as f:
            f.write("=" * 50 + "\n")
            f.write("PHISHING EMAIL ANALYSIS REPORT\n")
            f.write("=" * 50 + "\n")
            f.write(f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            # Email analysis results
            if 'email_analysis' in results:
                f.write("EMAIL ANALYSIS:\n")
                f.write("-" * 20 + "\n")
                email_result = results['email_analysis']
                if 'error' not in email_result:
                    f.write(f"Classification: {email_result.get('class', 'N/A')}\n")
                    f.write(f"Confidence: {email_result.get('probability', 0)*100:.2f}%\n")
                    f.write(f"Probabilities: {email_result.get('probabilities', {})}\n\n")
                else:
                    f.write(f"Error: {email_result['error']}\n\n")
            # URL analysis results
            if 'url_analysis' in results:
                f.write("URL ANALYSIS:\n")
                f.write("-" * 20 + "\n")
                for i, url_result in enumerate(results['url_analysis'], 1):
                    f.write(f"URL {i}: {url_result.get('url', 'N/A')}\n")
                    if 'error' not in url_result:
                        f.write(f"  Classification: {url_result.get('classification', 'N/A')}\n")
                        f.write(f"  Malicious detections: {url_result.get('malicious', 0)}\n")
                        f.write(f"  Suspicious detections: {url_result.get('suspicious', 0)}\n")
                        f.write(f"  Total scans: {url_result.get('total_scans', 0)}\n")
                    else:
                        f.write(f"  Error: {url_result['error']}\n")
                    f.write("\n")
            # Screenshot info
            if hasattr(self, 'screenshot_paths') and self.screenshot_paths:
                f.write("SCREENSHOTS:\n")
                f.write("-" * 20 + "\n")
                for path in self.screenshot_paths:
                    f.write(f"Saved: {path}\n")
                f.write("\n")
            # Original email content
            if 'email_content' in results:
                f.write("EMAIL CONTENT (PREVIEW):\n")
                f.write("-" * 20 + "\n")
                content = results['email_content'][:500] + "..." if len(results['email_content']) > 500 else results['email_content']
                f.write(f"{content}\n\n")
        self.content_log_path = log_file
        messagebox.showinfo("Log Saved", f"Log saved to {log_file}")

    def save_content_report(self):
        if not self.content_analysis_results:
            return
        reports_dir = "reports"
        os.makedirs(reports_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_file = os.path.join(reports_dir, f"email_content_report_{timestamp}.docx")
        results = self.content_analysis_results
        email_analysis = results.get('email_analysis', {})
        url_analysis = results.get('url_analysis', [])
        # Calculate summary statistics
        summary = {
            "email_classification": email_analysis.get('class', 'Unknown'),
            "email_confidence": email_analysis.get('probability', 0) * 100,
            "urls_found": len(url_analysis),
            "malicious_urls": len([u for u in url_analysis if u.get('classification') == 'Malicious']),
            "suspicious_urls": len([u for u in url_analysis if u.get('classification') == 'Suspicious']),
            "legitimate_urls": len([u for u in url_analysis if u.get('classification') == 'Legitimate']),
            "failed_analysis": len([u for u in url_analysis if 'error' in u])
        }
        # Determine overall risk level
        if summary['email_classification'].lower() == 'phishing' or summary['malicious_urls'] > 0:
            risk_level = "HIGH RISK"
        elif summary['suspicious_urls'] > 0 or summary['email_confidence'] > 70:
            risk_level = "MEDIUM RISK"
        else:
            risk_level = "LOW RISK"
        doc = Document()
        title = doc.add_heading('EMAIL CONTENTS ANALYSIS REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Analysis ID: {timestamp}")
        doc.add_paragraph()
        doc.add_heading('EXECUTIVE SUMMARY', level=1)
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Table Grid'
        summary_data = [
            ('Overall Risk Level:', risk_level),
            ('Email Classification:', summary['email_classification'].upper()),
            ('Confidence Level:', f"{summary['email_confidence']:.1f}%"),
            ('URLs Analyzed:', str(summary['urls_found']))
        ]
        for i, (key, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = key
            summary_table.cell(i, 1).text = str(value)
        doc.add_paragraph()
        # Email Analysis Section
        doc.add_heading('EMAIL ANALYSIS', level=1)
        email_table = doc.add_table(rows=2, cols=2)
        email_table.style = 'Table Grid'
        email_data = [
            ('Classification:', summary['email_classification']),
            ('Machine Learning Confidence:', f"{summary['email_confidence']:.1f}%")
        ]
        for i, (key, value) in enumerate(email_data):
            email_table.cell(i, 0).text = key
            email_table.cell(i, 1).text = str(value)
        doc.add_paragraph()
        # Add probability distribution
        if 'probabilities' in email_analysis:
            doc.add_paragraph('Probability Distribution:')
            prob_table = doc.add_table(rows=len(email_analysis['probabilities']), cols=2)
            prob_table.style = 'Table Grid'
            for i, (label, prob) in enumerate(email_analysis['probabilities'].items()):
                prob_table.cell(i, 0).text = label
                prob_table.cell(i, 1).text = f"{prob*100:.1f}%"
        doc.add_paragraph()
        # URL Analysis Summary
        doc.add_heading('URL ANALYSIS SUMMARY', level=1)
        url_summary_table = doc.add_table(rows=5, cols=2)
        url_summary_table.style = 'Table Grid'
        url_summary_data = [
            ('Total URLs Found:', str(summary['urls_found'])),
            ('Malicious URLs:', str(summary['malicious_urls'])),
            ('Suspicious URLs:', str(summary['suspicious_urls'])),
            ('Legitimate URLs:', str(summary['legitimate_urls'])),
            ('Analysis Failures:', str(summary['failed_analysis']))
        ]
        for i, (key, value) in enumerate(url_summary_data):
            url_summary_table.cell(i, 0).text = key
            url_summary_table.cell(i, 1).text = str(value)
        doc.add_paragraph()
        # Detailed URL Analysis
        if url_analysis:
            doc.add_heading('DETAILED URL ANALYSIS', level=1)
            url_detail_table = doc.add_table(rows=1, cols=6)
            url_detail_table.style = 'Table Grid'
            headers = ['#', 'URL', 'Classification', 'Malicious', 'Suspicious', 'Total Scans']
            for i, header in enumerate(headers):
                url_detail_table.cell(0, i).text = header
            for i, url_result in enumerate(url_analysis, 1):
                row = url_detail_table.add_row()
                if 'error' not in url_result:
                    row.cells[0].text = str(i)
                    row.cells[1].text = url_result.get('url', 'N/A')[:50] + "..." if len(url_result.get('url', '')) > 50 else url_result.get('url', 'N/A')
                    row.cells[2].text = url_result.get('classification', 'Unknown')
                    row.cells[3].text = str(url_result.get('malicious', 0))
                    row.cells[4].text = str(url_result.get('suspicious', 0))
                    row.cells[5].text = str(url_result.get('total_scans', 0))
                else:
                    row.cells[0].text = str(i)
                    row.cells[1].text = url_result.get('url', 'Unknown URL')[:50] + "..." if len(url_result.get('url', '')) > 50 else url_result.get('url', 'Unknown URL')
                    row.cells[2].text = "Error"
                    row.cells[3].text = "N/A"
                    row.cells[4].text = "N/A"
                    row.cells[5].text = url_result.get('error', 'Analysis failed')[:30] + "..." if len(url_result.get('error', '')) > 30 else url_result.get('error', 'Analysis failed')
        doc.add_paragraph()
        # Risk Assessment
        doc.add_heading('RISK ASSESSMENT', level=1)
        risk_table = doc.add_table(rows=1, cols=2)
        risk_table.style = 'Table Grid'
        risk_table.cell(0, 0).text = 'Risk Level:'
        risk_table.cell(0, 1).text = risk_level
        doc.add_paragraph('Risk Factors:')
        # Add risk factors
        risk_factors = []
        if summary['email_classification'].lower() == 'phishing':
            risk_factors.append("Email classified as PHISHING with high confidence")
        if summary['malicious_urls'] > 0:
            risk_factors.append(f"{summary['malicious_urls']} malicious URL(s) detected")
        if summary['suspicious_urls'] > 0:
            risk_factors.append(f"{summary['suspicious_urls']} suspicious URL(s) found")
        if summary['email_confidence'] > 80:
            risk_factors.append(f"High confidence level ({summary['email_confidence']:.1f}%) in classification")
        if not risk_factors:
            risk_factors.append("No significant risk factors identified")
        for factor in risk_factors:
            doc.add_paragraph(f" {factor}", style='List Bullet')
        doc.add_paragraph()
        # Recommendations
        doc.add_heading('RECOMMENDATIONS', level=1)
        if risk_level == "HIGH RISK":
            doc.add_paragraph('IMMEDIATE ACTION REQUIRED:', style='Heading 2')
            recommendations = [
                "Do NOT click any links in this email",
                "Do NOT download any attachments",
                "Do NOT provide any personal information",
                "Report this email to your IT security team",
                "Delete the email immediately",
                "Run a full antivirus scan if any links were clicked"
            ]
        elif risk_level == "MEDIUM RISK":
            doc.add_paragraph('EXERCISE CAUTION:', style='Heading 2')
            recommendations = [
                "Verify sender identity through alternative means",
                "Do not click suspicious links",
                "Hover over links to check destinations before clicking",
                "Be cautious with any requests for sensitive information",
                "Consider reporting to security team for further analysis"
            ]
        else:
            doc.add_paragraph('LOW RISK - STANDARD PRECAUTIONS:', style='Heading 2')
            recommendations = [
                "Email appears legitimate but remain vigilant",
                "Still verify sender if requesting sensitive information",
                "Keep security software updated",
                "Follow standard email security practices"
            ]
        for rec in recommendations:
            doc.add_paragraph(f" {rec}", style='List Bullet')
        doc.add_paragraph()
        # Next Steps
        doc.add_heading('NEXT STEPS', level=1)
        next_steps = [
            "Review the detailed findings above",
            "Follow the recommended actions based on risk level",
            "Consider additional security measures if high risk",
            "Document and report incidents according to company policy",
            "Monitor for similar threats in the future"
        ]
        for step in next_steps:
            doc.add_paragraph(step, style='List Number')
        doc.add_paragraph()
        footer = doc.add_paragraph('End of Report')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(report_file)
        self.content_report_path = report_file
        messagebox.showinfo("Report Saved", f"Report saved to {report_file}")

    # Header Analysis Tab Methods
    def load_header_file(self):
        path = filedialog.askopenfilename(filetypes=[("EML Files", "*.eml"), ("Text Files", "*.txt"), ("All Files", "*.*")])
        if path:
            self.header_file_path.set(path)

    def run_header_analysis(self):
        path = self.header_file_path.get()
        if not path:
            messagebox.showwarning("Input Required", "Please select a header file.")
            return
        self.header_results.config(state='normal')
        self.header_results.delete('1.0', tk.END)
        self.header_results.insert(tk.END, "Analyzing...\n")
        self.header_results.config(state='disabled')
        def worker():
            meta = analyze_headers_from_file(path)
            self.header_analysis_results = meta
            self.display_header_results(meta)
        threading.Thread(target=worker).start()

    def display_header_results(self, meta):
        self.header_results.config(state='normal')
        self.header_results.delete('1.0', tk.END)
        if 'error' in meta:
            messagebox.showwarning("Warning", "Please select .eml header file.")
            return
        else:
            self.header_results.insert(tk.END, f"Message ID: {meta['message-id']}\n")
            self.header_results.insert(tk.END, f"From: {meta['from']}\n")
            self.header_results.insert(tk.END, f"Date: {meta['dt']}\n")
            self.header_results.insert(tk.END, f"SPF: {'PASS' if meta['spf-record'] else 'FAIL'}\n")
            self.header_results.insert(tk.END, f"DKIM: {'PASS' if meta['dkim-record'] else 'FAIL'}\n")
            self.header_results.insert(tk.END, f"DMARC: {'PASS' if meta['dmarc-record'] else 'FAIL'}\n")
            self.header_results.insert(tk.END, f"IP Address: {meta['ip-address']}\n")
            self.header_results.insert(tk.END, f"Content-Type: {meta['content-type']}\n")
            self.header_results.insert(tk.END, f"Received Headers Count: {meta['multiple-received']}\n")
            if meta['reply-to']:
                self.header_results.insert(tk.END, f"Reply-To: {meta['reply-to']}\n")
            self.header_results.insert(tk.END, f"Threat Level: {meta['threat-level']}\n")
            self.header_results.insert(tk.END, f"Threat Score: {meta['threat-score']}/100\n")
            if meta.get('threat-reasons'):
                self.header_results.insert(tk.END, "Risk Factors:\n")
                for reason in meta['threat-reasons']:
                    self.header_results.insert(tk.END, f"  - {reason}\n")
            rec = get_header_recommendation(meta['threat-level'], meta['threat-score'], meta.get('threat-reasons', []))
            self.header_results.insert(tk.END, f"\nRECOMMENDATION: {rec}\n")
        self.header_results.config(state='disabled')
        self.btn_save_header_log.config(state='normal')
        self.btn_save_header_report.config(state='normal')

    def save_header_log(self):
        if not self.header_analysis_results:
            return
        logs_dir = "logs"
        os.makedirs(logs_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        log_file = os.path.join(logs_dir, f"email_headers_log_{timestamp}.txt")
        meta = self.header_analysis_results
        with open(log_file, 'w', encoding='utf-8') as log_file_obj:
            log_file_obj.write("EMAIL HEADER ANALYSIS LOG\n")
            log_file_obj.write("="*50 + "\n")
            log_file_obj.write(f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            log_file_obj.write(f"Message ID: {meta.get('message-id','')}\n")
            log_file_obj.write(f"From: {meta.get('from','')}\n")
            log_file_obj.write(f"Date: {meta.get('dt','')}\n")
            log_file_obj.write(f"Threat Level: {meta.get('threat-level','')}\n")
            log_file_obj.write(f"Threat Score: {meta.get('threat-score',0)}/100\n")
            log_file_obj.write(f"SPF: {'PASS' if meta.get('spf-record') else 'FAIL'}\n")
            log_file_obj.write(f"DKIM: {'PASS' if meta.get('dkim-record') else 'FAIL'}\n")
            log_file_obj.write(f"DMARC: {'PASS' if meta.get('dmarc-record') else 'FAIL'}\n")
            log_file_obj.write(f"IP Address: {meta.get('ip-address','')}\n")
            if meta.get('threat-reasons'):
                log_file_obj.write("\nRisk Factors:\n")
                for reason in meta['threat-reasons']:
                    log_file_obj.write(f"- {reason}\n")
            rec = get_header_recommendation(meta.get('threat-level',''), meta.get('threat-score',0), meta.get('threat-reasons',[]))
            log_file_obj.write(f"\nRecommendation: {rec}\n")
        self.header_log_path = log_file
        messagebox.showinfo("Log Saved", f"Log saved to {log_file}")

    def save_header_report(self):
        if not self.header_analysis_results:
            return
        reports_dir = "reports"
        os.makedirs(reports_dir, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_file = os.path.join(reports_dir, f"email_headers_report_{timestamp}.docx")
        meta = self.header_analysis_results
        doc = Document()
        title = doc.add_heading('EMAIL HEADERS ANALYSIS REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Analysis ID: {timestamp}")
        doc.add_paragraph()
        doc.add_heading('EXECUTIVE SUMMARY', level=1)
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Table Grid'
        summary_data = [
            ('Threat Level:', meta.get('threat-level','')),
            ('Risk Score:', f"{meta.get('threat-score',0)}/100"),
            ('Analysis Date:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ('Message ID:', meta.get('message-id',''))
        ]
        for i, (key, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = key
            summary_table.cell(i, 1).text = str(value)
        doc.add_paragraph()
        doc.add_heading('EMAIL DETAILS', level=1)
        details_table = doc.add_table(rows=4 + (1 if meta.get('reply-to') else 0), cols=2)
        details_table.style = 'Table Grid'
        details_data = [
            ('Message ID:', meta.get('message-id','')),
            ('From:', meta.get('from','')),
            ('Date Sent:', meta.get('dt','')),
            ('Content Type:', meta.get('content-type',''))
        ]
        if meta.get('reply-to'):
            details_data.append(('Reply-To:', meta.get('reply-to','')))
        for i, (key, value) in enumerate(details_data):
            details_table.cell(i, 0).text = key
            details_table.cell(i, 1).text = str(value)
        doc.add_paragraph()
        doc.add_heading('TECHNICAL ANALYSIS', level=1)
        tech_table = doc.add_table(rows=3, cols=2)
        tech_table.style = 'Table Grid'
        tech_data = [
            ('Sender IP Address:', meta.get('ip-address','')),
            ('Email Routing Hops:', str(meta.get('multiple-received',''))),
            ('Sender Client:', meta.get('sender-client','')[:100] + "..." if len(meta.get('sender-client','')) > 100 else meta.get('sender-client',''))
        ]
        for i, (key, value) in enumerate(tech_data):
            tech_table.cell(i, 0).text = key
            tech_table.cell(i, 1).text = str(value)
        doc.add_paragraph()
        doc.add_heading('AUTHENTICATION VERIFICATION', level=1)
        auth_table = doc.add_table(rows=3, cols=2)
        auth_table.style = 'Table Grid'
        auth_data = [
            ('SPF (Sender Policy Framework):', 'PASS' if meta.get('spf-record') else 'FAIL'),
            ('DKIM (DomainKeys Identified Mail):', 'PASS' if meta.get('dkim-record') else 'FAIL'),
            ('DMARC (Domain-based Message Authentication):', 'PASS' if meta.get('dmarc-record') else 'FAIL')
        ]
        for i, (key, value) in enumerate(auth_data):
            auth_table.cell(i, 0).text = key
            auth_table.cell(i, 1).text = str(value)
        doc.add_paragraph()
        doc.add_heading('RISK ASSESSMENT', level=1)
        if meta.get('threat-reasons'):
            doc.add_paragraph('Identified Risk Factors:')
            for i, reason in enumerate(meta['threat-reasons'], 1):
                doc.add_paragraph(f"{i}. {reason}", style='List Number')
        else:
            doc.add_paragraph('No significant risk factors identified.')
        doc.add_paragraph(f"Overall Risk Score: {meta.get('threat-score',0)}/100")
        doc.add_paragraph('Risk Level Interpretation:')
        doc.add_paragraph(' 0-29: LEGITIMATE (Low Risk)', style='List Bullet')
        doc.add_paragraph(' 30-69: SUSPICIOUS (Medium Risk)', style='List Bullet')
        doc.add_paragraph(' 70-100: MALICIOUS (High Risk)', style='List Bullet')
        doc.add_paragraph()
        doc.add_heading('RECOMMENDATIONS', level=1)
        rec = get_header_recommendation(meta.get('threat-level', ''), meta.get('threat-score', 0), meta.get('threat-reasons', []))
        doc.add_paragraph(f"Primary Recommendation: {rec}")
        doc.add_paragraph()
        doc.add_paragraph('General Security Best Practices:')
        best_practices = [
            'Verify sender identity through alternative communication channels',
            'Do not click suspicious links or download unexpected attachments',
            'Report suspicious emails to your IT security team',
            'Keep email security software updated',
            'Enable multi-factor authentication where possible'
        ]
        for practice in best_practices:
            doc.add_paragraph(practice, style='List Number')
        doc.add_paragraph()
        doc.add_heading('CONCLUSION', level=1)
        if meta.get('threat-level') == "MALICIOUS":
            conclusion = "This email exhibits multiple characteristics of malicious communication and should be treated as a security threat."
        elif meta.get('threat-level') == "SUSPICIOUS":
            conclusion = "This email shows suspicious characteristics that warrant careful review before any interaction."
        else:
            conclusion = "This email appears to be legitimate based on standard authentication and content analysis."
        conclusion += " Users should follow the recommended actions outlined above."
        doc.add_paragraph(conclusion)
        doc.add_paragraph()
        footer = doc.add_paragraph('End of Report')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(report_file)
        self.header_report_path = report_file
        messagebox.showinfo("Report Saved", f"Report saved to {report_file}")

if __name__ == "__main__":
    app = PhishingGUI()
    app.mainloop() 