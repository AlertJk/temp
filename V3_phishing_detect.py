import joblib
import numpy as np
import warnings
import re
import os
import requests
import time
from datetime import datetime
from urllib.parse import urlparse
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException
from sklearn.exceptions import InconsistentVersionWarning
import hashlib
import base64
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


warnings.filterwarnings("ignore", category=InconsistentVersionWarning)

class PhishingDetector:
    def __init__(self):
        self.model_path = 'D:/SEM 6/FYP/tools/phishing_model_v2.joblib'
        self.components = None
        self.vt_api_key = None
        self.load_model()
        self.load_api_keys()
    
    def load_model(self):
        try:
            if not os.path.exists(self.model_path):
                raise FileNotFoundError(f"Model file not found: {self.model_path}")
            self.components = joblib.load(self.model_path)
            print("Model loaded successfully")
        except Exception as e:
            print(f"Error loading model: {e}")
            raise
    
    def load_api_keys(self):
        try:
            self.vt_api_key = os.getenv('VIRUSTOTAL_API_KEY')
            
            if not self.vt_api_key:
                print("Warning: VirusTotal API key not found in environment variables")
                
        except Exception as e:
            print(f"Error loading API keys: {e}")
    
    def validate_input(self, text):        
        if len(text.strip()) == 0:
            raise ValueError("Input cannot be empty")
        
        if len(text) > 10000:
            raise ValueError("Input text is too long (max 10,000 characters)")
        
        return text.strip()
    
    def preprocess_text(self, text):
        text = re.sub(r'http\S+', '', text)
        text = re.sub(r'[^\w\s]', '', text)
        text = text.lower()
        text = re.sub(r'\s+', ' ', text).strip()
        return text
    
    def extract_urls(self, text):
        url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
        urls = re.findall(url_pattern, text)
        return list(set(urls))  # Remove duplicates
    
    def validate_url(self, url):
        try:
            parsed = urlparse(url)
            if not parsed.scheme or not parsed.netloc:
                return False
            
            if parsed.scheme not in ['http', 'https']:
                return False
            
            suspicious_patterns = [
                r'[0-9]{1,3}\.[0-9]{1,3}\.[0-9]{1,3}\.[0-9]{1,3}',
                r'bit\.ly|tinyurl|t\.co'
            ]
            
            for pattern in suspicious_patterns:
                if re.search(pattern, url):
                    print(f"Warning: URL contains suspicious pattern: {pattern}")
            
            return True
        except Exception:
            return False
    
    def analyze_url_virustotal(self, url):
        if not self.vt_api_key:
            return {"error": "VirusTotal API key not available"}
        
        try:
            url_id = base64.urlsafe_b64encode(url.encode()).decode().strip("=")
            
            headers = {
                "X-Apikey": self.vt_api_key,
                "Content-Type": "application/json"
            }
            
            response = requests.get(
                f"https://www.virustotal.com/api/v3/urls/{url_id}",
                headers=headers,
                timeout=30
            )
            
            if response.status_code == 404:
                submit_response = requests.post(
                    "https://www.virustotal.com/api/v3/urls",
                    headers=headers,
                    data={"url": url},
                    timeout=30
                )
                
                if submit_response.status_code == 200:
                    print("URL submitted for analysis. Waiting for results...")
                    time.sleep(15) 
                    
                    # Try to get results again
                    response = requests.get(
                        f"https://www.virustotal.com/api/v3/urls/{url_id}",
                        headers=headers,
                        timeout=30
                    )
            
            if response.status_code == 200:
                data = response.json()
                stats = data.get('data', {}).get('attributes', {}).get('last_analysis_stats', {})
                
                malicious = stats.get('malicious', 0)
                suspicious = stats.get('suspicious', 0)
                clean = stats.get('harmless', 0)
                undetected = stats.get('undetected', 0)
                
                total_scans = malicious + suspicious + clean + undetected
                
                if total_scans == 0:
                    return {"status": "No analysis data available", "url": url}
                
                # Determine classification
                if malicious > 0:
                    classification = "Malicious"
                elif suspicious > 2:  # More than 2 engines flagged as suspicious
                    classification = "Suspicious"
                else:
                    classification = "Legitimate"
                
                return {
                    "url": url,
                    "classification": classification,
                    "malicious": malicious,
                    "suspicious": suspicious,
                    "clean": clean,
                    "total_scans": total_scans,
                    "scan_date": data.get('data', {}).get('attributes', {}).get('last_analysis_date')
                }
            else:
                return {"error": f"API request failed with status {response.status_code}"}
                
        except requests.exceptions.Timeout:
            return {"error": "Request timeout"}
        except requests.exceptions.RequestException as e:
            return {"error": f"Request failed: {str(e)}"}
        except Exception as e:
            return {"error": f"Analysis failed: {str(e)}"}
    
    def capture_screenshot(self, url):
        try:
            if not self.validate_url(url):
                return {"error": "Invalid URL format"}
            
            chrome_options = Options()
            chrome_options.add_argument('--headless')
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--disable-gpu')
            chrome_options.add_argument('--window-size=1920,1080')
            chrome_options.add_argument('--disable-extensions')
            chrome_options.add_argument('--disable-plugins')
            chrome_options.add_argument('--disable-images')  # For faster loading
            chrome_options.add_argument('--disable-javascript')  # Security measure
            chrome_options.add_argument('--log-level=3')  # Suppress INFO, WARNING, ERROR
            chrome_options.add_argument('--silent')
            chrome_options.add_argument("--enable-unsafe-swiftshader")

            driver = webdriver.Chrome(options=chrome_options)
            driver.set_page_load_timeout(30)
            
            try:
                driver.get(url)
                
                # Wait for page to load
                WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.TAG_NAME, "body"))
                )
                
                # Create screenshots directory if it doesn't exist
                screenshot_dir = "screenshots"
                os.makedirs(screenshot_dir, exist_ok=True)
                
                # Generate filename based on URL and timestamp
                url_hash = hashlib.md5(url.encode()).hexdigest()[:8]
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"screenshot_{url_hash}_{timestamp}.png"
                filepath = os.path.join(screenshot_dir, filename)
                
                # Take screenshot
                driver.save_screenshot(filepath)
                print(f"Screenshot saved: {filepath}")
                
                return {"success": True, "filepath": filepath}
                
            except TimeoutException:
                return {"error": "Page load timeout"}
            except WebDriverException as e:
                return {"error": f"WebDriver error: {str(e)}"}
            finally:
                driver.quit()
                
        except Exception as e:
            return {"error": f"Screenshot capture failed: {str(e)}"}
    
    def predict_email(self, email_text):
        try:
            email_text = self.validate_input(email_text)
            
            tfidf_vectorizer = self.components['tfidf_vectorizer']
            voting_classifier = self.components['voting_classifier']
            label_encoder = self.components['label_encoder']
            
            # Preprocess
            processed_text = self.preprocess_text(email_text)
            email_tfidf = tfidf_vectorizer.transform([processed_text])
            
            # Predict
            prediction = voting_classifier.predict(email_tfidf)
            probabilities = voting_classifier.predict_proba(email_tfidf)
            
            return {
                'class': label_encoder.inverse_transform(prediction)[0],
                'probability': float(np.max(probabilities)),
                'probabilities': {
                    label_encoder.classes_[0]: float(probabilities[0][0]),
                    label_encoder.classes_[1]: float(probabilities[0][1])
                }
            }
        except Exception as e:
            return {"error": f"Prediction failed: {str(e)}"}
    
    def log_results(self, results):
        try:
            log_dir = "logs"
            os.makedirs(log_dir, exist_ok=True)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            log_file = os.path.join(log_dir, f"phishing_analysis_{timestamp}.txt")
            
            with open(log_file, 'w', encoding='utf-8') as f:
                f.write("=" * 50 + "\n")
                f.write("PHISHING EMAIL ANALYSIS REPORT\n")
                f.write("=" * 50 + "\n")
                f.write(f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                
                # Email analysis results
                if 'email_analysis' in results:
                    f.write("EMAIL ANALYSIS:\n")
                    f.write("-" * 20 + "\n")
                    email_result = results['email_analysis']
                    if 'error' not in email_result:
                        f.write(f"Classification: {email_result['class']}\n")
                        f.write(f"Confidence: {email_result['probability']*100:.2f}%\n")
                        f.write(f"Probabilities: {email_result['probabilities']}\n\n")
                    else:
                        f.write(f"Error: {email_result['error']}\n\n")
                
                # URL analysis results
                if 'url_analysis' in results:
                    f.write("URL ANALYSIS:\n")
                    f.write("-" * 20 + "\n")
                    for i, url_result in enumerate(results['url_analysis'], 1):
                        f.write(f"URL {i}: {url_result.get('url', 'N/A')}\n")
                        if 'error' not in url_result:
                            f.write(f"  Classification: {url_result.get('classification', 'N/A')}\n")
                            f.write(f"  Malicious detections: {url_result.get('malicious', 0)}\n")
                            f.write(f"  Suspicious detections: {url_result.get('suspicious', 0)}\n")
                            f.write(f"  Total scans: {url_result.get('total_scans', 0)}\n")
                        else:
                            f.write(f"  Error: {url_result['error']}\n")
                        f.write("\n")
                
                # Screenshot info
                if 'screenshots' in results:
                    f.write("SCREENSHOTS:\n")
                    f.write("-" * 20 + "\n")
                    for screenshot in results['screenshots']:
                        if 'error' not in screenshot:
                            f.write(f"Saved: {screenshot['filepath']}\n")
                        else:
                            f.write(f"Error: {screenshot['error']}\n")
                    f.write("\n")
                
                # Original email content
                if 'email_content' in results:
                    f.write("EMAIL CONTENT (PREVIEW):\n")
                    f.write("-" * 20 + "\n")
                    content = results['email_content'][:500] + "..." if len(results['email_content']) > 500 else results['email_content']
                    f.write(f"{content}\n\n")
            
            print(f"Results logged to: {log_file}")
            return log_file
            
        except Exception as e:
            print(f"Error logging results: {e}")
            return None
    

    def generate_report_simple(self, results):
        try:
            # Prepare data for report
            email_analysis = results.get('email_analysis', {})
            url_analysis = results.get('url_analysis', [])
            
            # Calculate summary statistics
            summary = {
                "email_classification": email_analysis.get('class', 'Unknown'),
                "email_confidence": email_analysis.get('probability', 0) * 100,
                "urls_found": len(url_analysis),
                "malicious_urls": len([u for u in url_analysis if u.get('classification') == 'Malicious']),
                "suspicious_urls": len([u for u in url_analysis if u.get('classification') == 'Suspicious']),
                "legitimate_urls": len([u for u in url_analysis if u.get('classification') == 'Legitimate']),
                "failed_analysis": len([u for u in url_analysis if 'error' in u])
            }
            
            # Determine overall risk level
            if summary['email_classification'].lower() == 'phishing' or summary['malicious_urls'] > 0:
                risk_level = "HIGH RISK"
            elif summary['suspicious_urls'] > 0 or summary['email_confidence'] > 70:
                risk_level = "MEDIUM RISK"
            else:
                risk_level = "LOW RISK"
            
            report_dir = "reports"
            os.makedirs(report_dir, exist_ok=True)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            report_file = os.path.join(report_dir, f"email_content_report_{timestamp}.docx")
            
            doc = Document()
            
            title = doc.add_heading('EMAIL CONTENTS ANALYSIS REPORT', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add timestamp and analysis ID
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Analysis ID: {timestamp}")
            doc.add_paragraph()
            
            doc.add_heading('EXECUTIVE SUMMARY', level=1)
            summary_table = doc.add_table(rows=4, cols=2)
            summary_table.style = 'Table Grid'
            
            summary_data = [
                ('Overall Risk Level:', risk_level),
                ('Email Classification:', summary['email_classification'].upper()),
                ('Confidence Level:', f"{summary['email_confidence']:.1f}%"),
                ('URLs Analyzed:', str(summary['urls_found']))
            ]
            
            for i, (key, value) in enumerate(summary_data):
                summary_table.cell(i, 0).text = key
                summary_table.cell(i, 1).text = str(value)
            
            doc.add_paragraph()
            
            # Email Analysis Section
            doc.add_heading('EMAIL ANALYSIS', level=1)
            email_table = doc.add_table(rows=2, cols=2)
            email_table.style = 'Table Grid'
            
            email_data = [
                ('Classification:', summary['email_classification']),
                ('Machine Learning Confidence:', f"{summary['email_confidence']:.1f}%")
            ]
            
            for i, (key, value) in enumerate(email_data):
                email_table.cell(i, 0).text = key
                email_table.cell(i, 1).text = str(value)

            doc.add_paragraph()
            # Add probability distribution
            if 'probabilities' in email_analysis:
                doc.add_paragraph('Probability Distribution:')
                prob_table = doc.add_table(rows=len(email_analysis['probabilities']), cols=2)
                prob_table.style = 'Table Grid'
                
                for i, (label, prob) in enumerate(email_analysis['probabilities'].items()):
                    prob_table.cell(i, 0).text = label
                    prob_table.cell(i, 1).text = f"{prob*100:.1f}%"
            
            doc.add_paragraph()
            
            # URL Analysis Summary
            doc.add_heading('URL ANALYSIS SUMMARY', level=1)
            url_summary_table = doc.add_table(rows=5, cols=2)
            url_summary_table.style = 'Table Grid'
            
            url_summary_data = [
                ('Total URLs Found:', str(summary['urls_found'])),
                ('Malicious URLs:', str(summary['malicious_urls'])),
                ('Suspicious URLs:', str(summary['suspicious_urls'])),
                ('Legitimate URLs:', str(summary['legitimate_urls'])),
                ('Analysis Failures:', str(summary['failed_analysis']))
            ]
            
            for i, (key, value) in enumerate(url_summary_data):
                url_summary_table.cell(i, 0).text = key
                url_summary_table.cell(i, 1).text = str(value)
            
            doc.add_paragraph()
            
            # Detailed URL Analysis
            if url_analysis:
                doc.add_heading('DETAILED URL ANALYSIS', level=1)
                
                url_detail_table = doc.add_table(rows=1, cols=6)
                url_detail_table.style = 'Table Grid'
                
                headers = ['#', 'URL', 'Classification', 'Malicious', 'Suspicious', 'Total Scans']
                for i, header in enumerate(headers):
                    url_detail_table.cell(0, i).text = header
                
                # Add URL data
                for i, url_result in enumerate(url_analysis, 1):
                    row = url_detail_table.add_row()
                    if 'error' not in url_result:
                        row.cells[0].text = str(i)
                        row.cells[1].text = url_result.get('url', 'N/A')[:50] + "..." if len(url_result.get('url', '')) > 50 else url_result.get('url', 'N/A')
                        row.cells[2].text = url_result.get('classification', 'Unknown')
                        row.cells[3].text = str(url_result.get('malicious', 0))
                        row.cells[4].text = str(url_result.get('suspicious', 0))
                        row.cells[5].text = str(url_result.get('total_scans', 0))
                    else:
                        row.cells[0].text = str(i)
                        row.cells[1].text = url_result.get('url', 'Unknown URL')[:50] + "..." if len(url_result.get('url', '')) > 50 else url_result.get('url', 'Unknown URL')
                        row.cells[2].text = "Error"
                        row.cells[3].text = "N/A"
                        row.cells[4].text = "N/A"
                        row.cells[5].text = url_result.get('error', 'Analysis failed')[:30] + "..." if len(url_result.get('error', '')) > 30 else url_result.get('error', 'Analysis failed')
            
            doc.add_paragraph()
            
            # Risk Assessment
            doc.add_heading('RISK ASSESSMENT', level=1)
            
            risk_table = doc.add_table(rows=1, cols=2)
            risk_table.style = 'Table Grid'
            risk_table.cell(0, 0).text = 'Risk Level:'
            risk_table.cell(0, 1).text = risk_level
            
            doc.add_paragraph('Risk Factors:')
            
            # Add risk factors
            risk_factors = []
            if summary['email_classification'].lower() == 'phishing':
                risk_factors.append("Email classified as PHISHING with high confidence")
            if summary['malicious_urls'] > 0:
                risk_factors.append(f"{summary['malicious_urls']} malicious URL(s) detected")
            if summary['suspicious_urls'] > 0:
                risk_factors.append(f"{summary['suspicious_urls']} suspicious URL(s) found")
            if summary['email_confidence'] > 80:
                risk_factors.append(f"High confidence level ({summary['email_confidence']:.1f}%) in classification")
            
            if not risk_factors:
                risk_factors.append("No significant risk factors identified")
            
            for factor in risk_factors:
                doc.add_paragraph(f" {factor}", style='List Bullet')
            
            doc.add_paragraph()
            
            # Recommendations
            doc.add_heading('RECOMMENDATIONS', level=1)
            
            if risk_level == "HIGH RISK":
                doc.add_paragraph('IMMEDIATE ACTION REQUIRED:', style='Heading 2')
                recommendations = [
                    "Do NOT click any links in this email",
                    "Do NOT download any attachments",
                    "Do NOT provide any personal information",
                    "Report this email to your IT security team",
                    "Delete the email immediately",
                    "Run a full antivirus scan if any links were clicked"
                ]
            elif risk_level == "MEDIUM RISK":
                doc.add_paragraph('EXERCISE CAUTION:', style='Heading 2')
                recommendations = [
                    "Verify sender identity through alternative means",
                    "Do not click suspicious links",
                    "Hover over links to check destinations before clicking",
                    "Be cautious with any requests for sensitive information",
                    "Consider reporting to security team for further analysis"
                ]
            else:
                doc.add_paragraph('LOW RISK - STANDARD PRECAUTIONS:', style='Heading 2')
                recommendations = [
                    "Email appears legitimate but remain vigilant",
                    "Still verify sender if requesting sensitive information",
                    "Keep security software updated",
                    "Follow standard email security practices"
                ]
            
            for rec in recommendations:
                doc.add_paragraph(f" {rec}", style='List Bullet')
            
            doc.add_paragraph()
            
            # Next Steps
            doc.add_heading('NEXT STEPS', level=1)
            next_steps = [
                "Review the detailed findings above",
                "Follow the recommended actions based on risk level",
                "Consider additional security measures if high risk",
                "Document and report incidents according to company policy",
                "Monitor for similar threats in the future"
            ]
            
            for step in next_steps:
                doc.add_paragraph(step, style='List Number')
            
            doc.add_paragraph()
            footer = doc.add_paragraph('End of Report')
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.save(report_file)
            
            print(f"Security Report generated: {report_file}")
            return {"success": True, "report": "Report generated successfully", "filepath": report_file}
            
        except Exception as e:
            return {"error": f"Report generation failed: {str(e)}"}
    
    def run_analysis(self):
        print("=" * 60)
        print("  PHISHING EMAIL DETECTION & URL ANALYSIS SYSTEM")
        print("=" * 60)
        
        try:
            print("\n1. Analyze Email Content")
            print("2. Analyze Email Headers")
            choice = input("\n Select analysis type (1 or 2): ").strip()

            if choice == "2":
                print("\n Redirecting to email header analysis...")
                
                from V2_headers_copy import analyze_email_headers
                analyze_email_headers()

            elif choice == "1":     
                email_text = input("\n Enter the email text to analyze: ").strip()
                
                if not email_text:
                    print(" No email text provided")
                    return
                
                results = {'email_content': email_text}
                
                # Analyze email
                print("\n Analyzing email content...")
                email_result = self.predict_email(email_text)
                results['email_analysis'] = email_result
                
                if 'error' in email_result:
                    print(f" Email analysis failed: {email_result['error']}")
                    return
                
                print(f" Email Classification: {email_result['class']}")
                print(f" Confidence: {email_result['probability']*100:.2f}%")
                
                # Extract and analyze URLs
                urls = self.extract_urls(email_text)
                results['url_analysis'] = []
                results['screenshots'] = []
                
                if urls:
                    print(f"\n Found {len(urls)} URL(s) in email:")
                    for i, url in enumerate(urls, 1):
                        print(f"   {i}. {url}")
                    
                    analyze_urls = input("\n Analyze URLs with VirusTotal? (y/n): ").lower().strip()
                    
                    if analyze_urls == 'y':
                        print("\n Analyzing URLs with VirusTotal...")
                        for url in urls:
                            if self.validate_url(url):
                                print(f"   Analyzing: {url}")
                                url_result = self.analyze_url_virustotal(url)
                                results['url_analysis'].append(url_result)
                                
                                if 'error' not in url_result:
                                    print(f"    Classification: {url_result.get('classification', 'Unknown')}")
                                    print(f"    Detections: {url_result.get('malicious', 0)} malicious, {url_result.get('suspicious', 0)} suspicious")
                                else:
                                    print(f"    Error: {url_result['error']}")
                            else:
                                print(f"    Skipping invalid URL: {url}")
                                results['url_analysis'].append({"url": url, "error": "Invalid URL"})
                    
                    # Screenshot capture
                    capture_screenshots = input("\n Capture screenshots of URLs? (y/n): ").lower().strip()
                    
                    if capture_screenshots == 'y':
                        print("\n Capturing screenshots...")
                        for url in urls:
                            if self.validate_url(url):
                                print(f"   Capturing: {url}")
                                screenshot_result = self.capture_screenshot(url)
                                results['screenshots'].append(screenshot_result)
                                
                                if 'error' in screenshot_result:
                                    print(f"    Error: {screenshot_result['error']}")
                            else:
                                results['screenshots'].append({"url": url, "error": "Invalid URL"})
                else:
                    print("\n No URLs found in email")
                
                # Log results
                log_results = input("\n Log results to file? (y/n): ").lower().strip()
                
                if log_results == 'y':
                    log_file = self.log_results(results)
                    if log_file:
                        results['log_file'] = log_file
                
                # SGenerate security report
                generate_report = input("\n Generate security report? (y/n): ").lower().strip()
                
                if generate_report == 'y':
                    print("\n Generating security report...")
                    report_result = self.generate_report_simple(results)
                    
                    if 'error' not in report_result:
                        print(report_result['report'])
                        # results['report'] = report_result
                    else:
                        print(f" Report generation failed: {report_result['error']}")
                
                print("\n Analysis completed successfully!")
            else:
                print(" Invalid choice. Please select 1 or 2.")
                
        except KeyboardInterrupt:
                print("\n\n Analysis interrupted by user")
        except Exception as e:
                print(f"\n Unexpected error: {e}")

def main():
    try:
        detector = PhishingDetector()
        detector.run_analysis()
    except Exception as e:
        print(f" System initialization failed: {e}")
        print("\nPlease ensure:")
        print("1. Model file exists at the specified path")
        print("2. Required packages are installed (selenium, requests, etc.)")
        print("3. API keys are set in environment variables")

if __name__ == "__main__":
    main()
