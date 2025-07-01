import email
import re
import os
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from rich.console import Console
from rich.table import Table
from rich.panel import Panel
from rich.text import Text
from rich.progress import Progress, SpinnerColumn, TextColumn

console = Console()

def analyze_email_headers():
    # Rich input prompt
    console.print("\n[bold bright_cyan]EMAIL HEADER ANALYSIS TOOL[/bold bright_cyan]")
    console.print("[dim bright_white]Enter the path to your email file for comprehensive analysis[/dim bright_white]")
    
    file_path = input("File path: ").strip().strip('"').strip("'")
    
    if not file_path:
        console.print("[bold red]No file path provided[/bold red]")
        return
    
    try:
        # Loading animation
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
        ) as progress:
            task = progress.add_task("[bright_cyan]Loading email file...", total=None)
            
            with open(file_path, 'r', encoding='utf-8') as f:
                msg = email.message_from_file(f)
            
            progress.update(task, description="[bright_cyan]Parsing email headers...")
            parser = email.parser.HeaderParser()
            headers = parser.parsestr(msg.as_string())

        meta = {
            "message-id": "",
            "spf-record": False,
            "dkim-record": False,
            "dmarc-record": False,
            "spoofed": False,
            "ip-address": "",
            "sender-client": "",
            "spoofed-mail": "",
            "dt": "",
            "content-type": "",
            "from": "",
            "reply-to": "",
            "return-path": "",
            "x-originating-ip": "",
            "multiple-received": 0,
            "threat-level": "LEGITIMATE",
            "threat-score": 0
        }

        # Count received headers to detect email bouncing
        received_count = 0
        
        console.print("\n[bright_yellow] Analyzing email headers...[/bright_yellow]")
        
        for h in headers.items():
            # Message ID
            if h[0].lower() == "message-id":
                meta["message-id"] = h[1]

            # From header
            if h[0].lower() == "from":
                meta["from"] = h[1]

            # Return path for comparison
            if h[0].lower() == "return-path":
                meta["return-path"] = h[1]

            # Count received headers
            if h[0].lower() == "received":
                received_count += 1
                if received_count == 1:  # Use first received header as sender client
                    meta["sender-client"] = h[1]

            # X-Originating-IP header
            if h[0].lower() == "x-originating-ip":
                meta["x-originating-ip"] = h[1]

            # Authentication detected by mail server
            if h[0].lower() == "authentication-results":
                if re.search(r"spf=pass", h[1], re.IGNORECASE):
                    meta["spf-record"] = True
                if re.search(r"dkim=pass", h[1], re.IGNORECASE):
                    meta["dkim-record"] = True
                if re.search(r"dmarc=pass", h[1], re.IGNORECASE):
                    meta["dmarc-record"] = True
                if re.search(r"does not designate|spf=fail|dkim=fail|dmarc=fail", h[1], re.IGNORECASE):
                    meta["spoofed"] = True
                if re.search(r"(\d{1,3}\.){3}\d{1,3}", h[1]):
                    ip = re.search(r"(\d{1,3}\.){3}\d{1,3}", h[1])
                    meta["ip-address"] = str(ip.group())

            if h[0].lower() == "reply-to":
                meta["reply-to"] = h[1]
                meta["spoofed-mail"] = h[1]

            if h[0].lower() == "date":
                meta["dt"] = h[1]

            if h[0].lower() == "content-type":
                meta["content-type"] = h[1]

        meta["multiple-received"] = received_count

        # Threat Assessment Logic
        threat_score = 0
        threat_reasons = []

        console.print("[bright_magenta] Performing threat assessment...[/bright_magenta]")

        # Authentication failures (high risk)
        if meta["spoofed"]:
            threat_score += 40
            threat_reasons.append("Authentication failures detected")

        if not meta["spf-record"]:
            threat_score += 15
            threat_reasons.append("SPF record failed")

        if not meta["dkim-record"]:
            threat_score += 15
            threat_reasons.append("DKIM verification failed")

        if not meta["dmarc-record"]:
            threat_score += 15
            threat_reasons.append("DMARC verification failed")

        # Mismatched sender information
        if meta["from"] and meta["reply-to"] and meta["from"] != meta["reply-to"]:
            threat_score += 20
            threat_reasons.append("From and Reply-To headers mismatch")

        # Excessive email hops (potential email bouncing)
        if meta["multiple-received"] > 8:
            threat_score += 15
            threat_reasons.append("Excessive email routing detected")

        # No proper sender identification
        if not meta["message-id"]:
            threat_score += 10
            threat_reasons.append("Missing Message-ID")

        # Determine threat level
        if threat_score >= 70:
            meta["threat-level"] = "MALICIOUS"
        elif threat_score >= 30:
            meta["threat-level"] = "SUSPICIOUS"
        else:
            meta["threat-level"] = "LEGITIMATE"

        meta["threat-score"] = threat_score
        meta["threat-reasons"] = threat_reasons

        # Display Rich Results
        console.print("\n")
        console.print("=" * 70)
        console.print("[bold bright_cyan on blue]|\t \t EMAIL HEADER ANALYSIS RESULTS \t \t|[/bold bright_cyan on blue]")
        console.print("=" * 70)

        # Basic Information Table
        basic_table = Table(title="[bold bright_white] Basic Email Information[/bold bright_white]", 
                           show_header=True, header_style="bold bright_cyan")
        basic_table.add_column("Field", style="bright_yellow", width=20)
        basic_table.add_column("Value", style="bright_white", width=50)
        
        basic_table.add_row("Message ID", meta['message-id'][:60] + "..." if len(meta['message-id']) > 60 else meta['message-id'])
        basic_table.add_row("From", meta['from'])
        basic_table.add_row("Date", meta['dt'])
        if meta['reply-to']:
            basic_table.add_row("Reply-To", meta['reply-to'])
        
        console.print(basic_table)
        console.print()

        # Authentication Status Table
        auth_table = Table(title="[bold bright_white]Authentication Status[/bold bright_white]", 
                          show_header=True, header_style="bold bright_cyan")
        auth_table.add_column("Protocol", style="bright_yellow", width=15)
        auth_table.add_column("Status", width=15)
        auth_table.add_column("Description", style="dim bright_white", width=40)
        
        # SPF Status
        spf_status = "[bold green]PASS[/bold green]" if meta['spf-record'] else "[bold red]FAIL[/bold red]"
        auth_table.add_row("SPF", spf_status, "Sender Policy Framework")
        
        # DKIM Status
        dkim_status = "[bold green]PASS[/bold green]" if meta['dkim-record'] else "[bold red]FAIL[/bold red]"
        auth_table.add_row("DKIM", dkim_status, "DomainKeys Identified Mail")
        
        # DMARC Status
        dmarc_status = "[bold green]PASS[/bold green]" if meta['dmarc-record'] else "[bold red]FAIL[/bold red]"
        auth_table.add_row("DMARC", dmarc_status, "Domain-based Message Authentication")
        
        console.print(auth_table)
        console.print()

        # Technical Details Table
        tech_table = Table(title="[bold bright_white]Technical Details[/bold bright_white]", 
                          show_header=True, header_style="bold bright_cyan")
        tech_table.add_column("Field", style="bright_yellow", width=20)
        tech_table.add_column("Value", style="bright_white", width=50)
        
        tech_table.add_row("IP Address", meta['ip-address'] if meta['ip-address'] else "[dim]Not detected[/dim]")
        tech_table.add_row("Content-Type", meta['content-type'])
        tech_table.add_row("Received Headers", str(meta['multiple-received']))
        
        console.print(tech_table)
        console.print()

        # Threat Assessment Panel
        threat_color = "red" if meta['threat-level'] == "MALICIOUS" else "yellow" if meta['threat-level'] == "SUSPICIOUS" else "green"
        
        threat_panel = Panel(
            f"[bold {threat_color}]THREAT LEVEL: {meta['threat-level']}[/bold {threat_color}]\n"
            f"[bold {threat_color}]Risk Score: {meta['threat-score']}/100[/bold {threat_color}]",
            title="[bold bright_red] THREAT ASSESSMENT[/bold bright_red]",
            border_style=threat_color,
            expand=False
        )
        console.print(threat_panel, justify="center")
        console.print()

        # Risk Factors
        if threat_reasons:
            console.print("[bold bright_red]IDENTIFIED RISK FACTORS:[/bold bright_red]")
            for i, reason in enumerate(threat_reasons, 1):
                console.print(f"   [bright_red]{i}.[/bright_red] [bright_white]{reason}[/bright_white]")
            console.print()

        # Get and display recommendation
        recommendation = get_recommendation(meta["threat-level"])
        
        rec_color = "red" if meta['threat-level'] == "MALICIOUS" else "yellow" if meta['threat-level'] == "SUSPICIOUS" else "green"
        rec_panel = Panel(
            recommendation,
            title=f"[bold {rec_color}]SECURITY RECOMMENDATION[/bold {rec_color}]",
            border_style=rec_color,
            expand=False
        )
        console.print(rec_panel, justify="center")
        console.print()

        console.print("=" * 80)

        # Options for saving results
        console.print("\n[bold bright_blue]SAVE OPTIONS[/bold bright_blue]")
        console.print("[bright_cyan]1.[/bright_cyan] [bright_white]Save to text log file[/bright_white]")
        console.print("[bright_cyan]2.[/bright_cyan] [bright_white]Generate detailed DOCX report[/bright_white]")
        console.print()
        while True:
            log_choice = input("Save results to text file? (y/n): ").strip().lower()
            if log_choice == 'y':
                console.print("[bright_blue]Saving to log file...[/bright_blue]")
                log_results(meta, threat_reasons, recommendation, file_path)
                break
            elif log_choice == 'n':
                console.print("[yellow] Skipping log saving[/yellow]")
                break
            else:
                console.print("[bold red]= Invalid choice. Please select 'y' or 'n'.[/bold red]")
                console.print("[dim bright_yellow] Tip: Enter either 'y' for log saving or 'n' to skip log saving[/dim bright_yellow]\n\n")


        while True:
            report_choice = input("\nGenerate detailed report document? (y/n): ").strip().lower()
            if report_choice == 'y':
                console.print("[bright_blue]Generating comprehensive report...[/bright_blue]")
                generate_report(meta, threat_reasons, recommendation, file_path)
                break
            elif log_choice == 'n':
                console.print("[yellow] Skipping report generation[/yellow]")
                break
            else:
                console.print("[bold red]= Invalid choice. Please select 'y' or 'n'.[/bold red]")
                console.print("[dim bright_yellow] Tip: Enter either 'y' for report generation or 'n' to skip report generation[/dim bright_yellow]\n\n")

        # Success message
        console.print("\n[bold green]" + "=" * 70 + "[/bold green]")
        console.print("[bold green on blue]|\t \t ANALYSIS COMPLETED SUCCESSFULLY! \t \t|[/bold green on blue]")
        console.print("[bold green]" + "=" * 70 + "[/bold green]")

    except FileNotFoundError:
        console.print(f"[bold red]File not found: {file_path}[/bold red]")
        console.print("[dim bright_yellow]Please check the file path and try again[/dim bright_yellow]")
    except Exception as e:
        console.print(f"[bold red]An error occurred: {str(e)}[/bold red]")
        console.print("[dim bright_red]Please check your file format and try again[/dim bright_red]")

def get_recommendation(threat_level):
    """Generate recommendation based on threat assessment"""
    if threat_level == "MALICIOUS":
        return "[bold red]DO NOT INTERACT[/bold red] with this email. Delete immediately and report to IT security team. This email shows strong indicators of malicious intent."
    elif threat_level == "SUSPICIOUS":
        return "[bold yellow] EXERCISE CAUTION[/bold yellow]. Do not click links or download attachments. Verify sender through alternative communication method before taking any action."
    else:
        return "[bold green]Email appears legitimate[/bold green]. However, always exercise standard email security practices."

def log_results(meta, threat_reasons, recommendation, original_file):
    """Log analysis results to a text file in logs folder"""
    # Create logs directory if it doesn't exist
    logs_dir = "logs"
    if not os.path.exists(logs_dir):
        os.makedirs(logs_dir)
        console.print(f"[dim bright_cyan]Created logs directory: {logs_dir}[/dim bright_cyan]")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_filename = os.path.join(logs_dir, f"email_headers_log_{timestamp}.txt")
    
    try:
        with open(log_filename, 'w', encoding='utf-8') as log_file:
            log_file.write("EMAIL HEADER ANALYSIS LOG\n")
            log_file.write("="*50 + "\n")
            log_file.write(f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            log_file.write(f"Original File: {original_file}\n")
            log_file.write(f"Message ID: {meta['message-id']}\n")
            log_file.write(f"From: {meta['from']}\n")
            log_file.write(f"Date: {meta['dt']}\n")
            log_file.write(f"Threat Level: {meta['threat-level']}\n")
            log_file.write(f"Threat Score: {meta['threat-score']}/100\n")
            log_file.write(f"SPF: {'PASS' if meta['spf-record'] else 'FAIL'}\n")
            log_file.write(f"DKIM: {'PASS' if meta['dkim-record'] else 'FAIL'}\n")
            log_file.write(f"DMARC: {'PASS' if meta['dmarc-record'] else 'FAIL'}\n")
            log_file.write(f"IP Address: {meta['ip-address']}\n")
            
            if threat_reasons:
                log_file.write("\nRisk Factors:\n")
                for reason in threat_reasons:
                    log_file.write(f"- {reason}\n")
            
            # Clean recommendation text for file
            clean_recommendation = recommendation.replace("[bold red]", "").replace("[/bold red]", "")
            clean_recommendation = clean_recommendation.replace("[bold yellow]", "").replace("[/bold yellow]", "")
            clean_recommendation = clean_recommendation.replace("[bold green]", "").replace("[/bold green]", "")
            
            log_file.write(f"\nRecommendation: {clean_recommendation}\n")
            
        console.print(f"[bold green]Results logged to: {log_filename}[/bold green]")
    except Exception as e:
        console.print(f"[bold red]Error creating log file: {str(e)}[/bold red]")

def generate_report(meta, threat_reasons, recommendation, original_file):
    """Generate a detailed report document as DOCX in reports folder"""
    # Create reports directory if it doesn't exist
    reports_dir = "reports"
    if not os.path.exists(reports_dir):
        os.makedirs(reports_dir)
        console.print(f"[dim bright_cyan]Created reports directory: {reports_dir}[/dim bright_cyan]")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_filename = os.path.join(reports_dir, f"email_headers_report_{timestamp}.docx")
    
    try:
        # Create a new Document
        doc = Document()
        
        # Title
        title = doc.add_heading('EMAIL HEADERS ANALYSIS REPORT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a line break
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading('EXECUTIVE SUMMARY', level=1)
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_data = [
            ('Threat Level:', meta['threat-level']),
            ('Risk Score:', f"{meta['threat-score']}/100"),
            ('Analysis Date:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ('Email File Analyzed:', original_file)
        ]
        
        for i, (key, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = key
            summary_table.cell(i, 1).text = str(value)
        
        doc.add_paragraph()
        
        # Email Details
        doc.add_heading(' EMAIL DETAILS', level=1)
        details_table = doc.add_table(rows=4 + (1 if meta['reply-to'] else 0), cols=2)
        details_table.style = 'Table Grid'
        
        details_data = [
            (' Message ID:', meta['message-id']),
            (' From:', meta['from']),
            (' Date Sent:', meta['dt']),
            (' Content Type:', meta['content-type'])
        ]
        
        if meta['reply-to']:
            details_data.append(('↩ Reply-To:', meta['reply-to']))
        
        for i, (key, value) in enumerate(details_data):
            details_table.cell(i, 0).text = key
            details_table.cell(i, 1).text = str(value)
        
        doc.add_paragraph()
        
        # Technical Analysis
        doc.add_heading(' TECHNICAL ANALYSIS', level=1)
        tech_table = doc.add_table(rows=3, cols=2)
        tech_table.style = 'Table Grid'
        
        tech_data = [
            (' Sender IP Address:', meta['ip-address']),
            (' Email Routing Hops:', str(meta['multiple-received'])),
            (' Sender Client:', meta['sender-client'][:100] + "..." if len(meta['sender-client']) > 100 else meta['sender-client'])
        ]
        
        for i, (key, value) in enumerate(tech_data):
            tech_table.cell(i, 0).text = key
            tech_table.cell(i, 1).text = str(value)
        
        doc.add_paragraph()
        
        # Authentication Results
        doc.add_heading(' AUTHENTICATION VERIFICATION', level=1)
        auth_table = doc.add_table(rows=3, cols=2)
        auth_table.style = 'Table Grid'
        
        auth_data = [
            (' SPF (Sender Policy Framework):', ' PASS' if meta['spf-record'] else ' FAIL'),
            (' DKIM (DomainKeys Identified Mail):', ' PASS' if meta['dkim-record'] else ' FAIL'),
            (' DMARC (Domain-based Message Authentication):', ' PASS' if meta['dmarc-record'] else ' FAIL')
        ]
        
        for i, (key, value) in enumerate(auth_data):
            auth_table.cell(i, 0).text = key
            auth_table.cell(i, 1).text = str(value)
        
        doc.add_paragraph()
        
        # Risk Assessment
        doc.add_heading(' RISK ASSESSMENT', level=1)
        
        if threat_reasons:
            doc.add_paragraph(' Identified Risk Factors:')
            for i, reason in enumerate(threat_reasons, 1):
                doc.add_paragraph(f"{i}. {reason}", style='List Number')
        else:
            doc.add_paragraph(' No significant risk factors identified.')
    
        doc.add_paragraph(f" Overall Risk Score: {meta['threat-score']}/100")
        
        doc.add_paragraph(' Risk Level Interpretation:')
        doc.add_paragraph(' 0-29: LEGITIMATE (Low Risk)', style='List Bullet')
        doc.add_paragraph(' 30-69: SUSPICIOUS (Medium Risk)', style='List Bullet')
        doc.add_paragraph(' 70-100: MALICIOUS (High Risk)', style='List Bullet')
        
        doc.add_paragraph()
        
        # Recommendations
        doc.add_heading(' RECOMMENDATIONS', level=1)
        # Clean recommendation text for document
        clean_recommendation = recommendation.replace("[bold red]", "").replace("[/bold red]", "")
        clean_recommendation = clean_recommendation.replace("[bold yellow]", "").replace("[/bold yellow]", "")
        clean_recommendation = clean_recommendation.replace("[bold green]", "").replace("[/bold green]", "")
        
        doc.add_paragraph(f" Primary Recommendation: {clean_recommendation}")
        doc.add_paragraph()
        
        doc.add_paragraph(' General Security Best Practices:')
        best_practices = [
            ' Verify sender identity through alternative communication channels',
            ' Do not click suspicious links or download unexpected attachments',
            ' Report suspicious emails to your IT security team',
            ' Keep email security software updated',
            ' Enable multi-factor authentication where possible'
        ]
        
        for practice in best_practices:
            doc.add_paragraph(practice, style='List Number')
        
        doc.add_paragraph()
        
        # Conclusion
        doc.add_heading(' CONCLUSION', level=1)
        if meta['threat-level'] == "MALICIOUS":
            conclusion = " This email exhibits multiple characteristics of malicious communication and should be treated as a security threat."
        elif meta['threat-level'] == "SUSPICIOUS":
            conclusion = " This email shows suspicious characteristics that warrant careful review before any interaction."
        else:
            conclusion = " This email appears to be legitimate based on standard authentication and content analysis."
        
        conclusion += " Users should follow the recommended actions outlined above."
        doc.add_paragraph(conclusion)
        
        # Add footer
        doc.add_paragraph()
        footer = doc.add_paragraph(' End of Report')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save the document
        doc.save(report_filename)
        console.print(f"[bold green] Detailed report generated: {report_filename}[/bold green]")
        
    except Exception as e:
        console.print(f"[bold red] Error generating report: {str(e)}[/bold red]")
        console.print("[dim bright_yellow] Note: Make sure you have python-docx installed: pip install python-docx[/dim bright_yellow]")